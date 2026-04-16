from __future__ import annotations

import os
import re
from pathlib import Path

from app.models import Report
from app.services.google_drive import GoogleDriveUploadService


class WordExportService:
    def __init__(self, exports_dir: Path):
        self.exports_dir = exports_dir
        self.exports_dir.mkdir(parents=True, exist_ok=True)
        self.upload_meta_dir = self.exports_dir / ".upload_meta"
        self.upload_meta_dir.mkdir(parents=True, exist_ok=True)
        self.google_drive = GoogleDriveUploadService()

    def _safe_slug(self, value: str) -> str:
        cleaned = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
        return cleaned[:80] or "report"

    def _docx_imports(self):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches, Pt, RGBColor

            return {
                "Document": Document,
                "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
                "OxmlElement": OxmlElement,
                "qn": qn,
                "Inches": Inches,
                "Pt": Pt,
                "RGBColor": RGBColor,
            }
        except Exception as exc:
            raise RuntimeError(
                "DOCX export requires python-docx. Install dependencies with `pip install -r requirements.txt`."
            ) from exc

    def _set_cell_shading(self, cell, fill: str, OxmlElement, qn) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _style_document(self, document, Pt, Inches) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        for style_name, size in (("Heading 1", 15), ("Heading 2", 12)):
            style = document.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = True

    def _add_title(self, document, report: Report, WD_ALIGN_PARAGRAPH, Pt, RGBColor) -> None:
        title = document.add_paragraph()
        run = title.add_run(f"{report.company_name} Tax Incentive Summary")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x11, 0x2D, 0x4E)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_paragraph(self, document, text: str | None) -> None:
        if text:
            p = document.add_paragraph(text.strip())
            p.paragraph_format.space_after = 6

    def _add_paragraph_break(self, document) -> None:
        p = document.add_paragraph("")
        p.paragraph_format.space_after = 6

    def _add_emphasis_paragraph(self, document, lead: str, emphasis: str, tail: str) -> None:
        p = document.add_paragraph()
        if lead:
            p.add_run(lead)
        if emphasis:
            r = p.add_run(emphasis)
            r.bold = True
        if tail:
            p.add_run(tail)
        p.paragraph_format.space_after = 6

    def _add_section_heading(self, document, text: str) -> None:
        if len(document.paragraphs) > 0:
            self._add_paragraph_break(document)
        h = document.add_paragraph(text, style="Heading 1")
        h.paragraph_format.space_before = 12
        h.paragraph_format.space_after = 6

    def _add_table(self, document, headers: list[str], rows: list[list[str]], OxmlElement, qn, RGBColor, Pt) -> None:
        self._add_paragraph_break(document)
        row_count = max(1, len(rows)) + 1
        table = document.add_table(rows=row_count, cols=len(headers))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header
            self._set_cell_shading(hdr_cells[idx], "1F4E78", OxmlElement, qn)
            p = hdr_cells[idx].paragraphs[0]
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)

        body_rows = rows or [["No data available"] + [""] * (len(headers) - 1)]
        for ridx, row in enumerate(body_rows, start=1):
            cells = table.rows[ridx].cells
            for cidx, value in enumerate(row):
                cells[cidx].text = str(value or "")
                p = cells[cidx].paragraphs[0]
                p.paragraph_format.space_before = 0
                p.paragraph_format.space_after = 0
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if ridx % 2 == 0:
                for cell in cells:
                    self._set_cell_shading(cell, "F4F7FB", OxmlElement, qn)
        self._add_paragraph_break(document)

    def _build_document(self, report: Report, include_confidence: bool = False):
        deps = self._docx_imports()
        Document = deps["Document"]
        WD_ALIGN_PARAGRAPH = deps["WD_ALIGN_PARAGRAPH"]
        OxmlElement = deps["OxmlElement"]
        qn = deps["qn"]
        Inches = deps["Inches"]
        Pt = deps["Pt"]
        RGBColor = deps["RGBColor"]

        template_path = os.getenv("WORD_TEMPLATE_PATH", "").strip()
        if template_path and Path(template_path).exists():
            document = Document(template_path)
        else:
            document = Document()

        self._style_document(document, Pt, Inches)
        self._add_title(document, report, WD_ALIGN_PARAGRAPH, Pt, RGBColor)

        self._add_section_heading(document, "Company Sector and Industry Analysis")
        self._add_paragraph(document, f"Company Description: {report.narrative.get('company_description', '')}")
        self._add_paragraph_break(document)
        self._add_paragraph(document, report.narrative.get("sector_summary", ""))

        self._add_section_heading(document, "GA Job Tax Credit")
        self._add_paragraph(document, report.narrative.get("ga_jtc_intro", ""))
        self._add_paragraph_break(document)
        self._add_paragraph(document, report.narrative.get("ga_jtc_note", ""))
        self._add_table(
            document,
            [
                "GA Location",
                "County",
                "County Tier",
                "Special Designation",
                "Job Creation Threshold",
                "Per Job Credit Amount",
            ],
            [
                [
                    loc.address,
                    loc.county or "-",
                    loc.ga_tier_label or "Unmapped",
                    loc.special_designation or "None",
                    loc.job_creation_threshold or "Unavailable",
                    loc.per_job_credit_amount or "Unavailable",
                ]
                for loc in report.locations
            ],
            OxmlElement,
            qn,
            RGBColor,
            Pt,
        )
        prior_years = [str(year) for year in report.narrative.get("ga_jtc_prior_years", [])]
        prior_rows = list(report.narrative.get("ga_jtc_prior_rows", []))
        if prior_years and prior_rows:
            self._add_table(
                document,
                ["Addresses Prior Year Tiers"] + prior_years,
                [
                    [str(row.get("address", ""))] + [str(tier) for tier in row.get("tiers", [])]
                    for row in prior_rows
                ],
                OxmlElement,
                qn,
                RGBColor,
                Pt,
            )
        else:
            self._add_paragraph(
                document,
                "Prior Year Tier History: No prior-year tier history available for the entered location(s).",
            )

        self._add_section_heading(document, "Georgia Retraining Tax Credit")
        intro_lead = str(report.narrative.get("retraining_intro_lead", ""))
        intro_emphasis = str(report.narrative.get("retraining_intro_emphasis", ""))
        intro_tail = str(report.narrative.get("retraining_intro_tail", ""))
        if intro_emphasis:
            self._add_emphasis_paragraph(document, intro_lead, intro_emphasis, intro_tail)
        else:
            self._add_paragraph(document, report.narrative.get("retraining_intro", ""))
        self._add_paragraph_break(document)
        self._add_paragraph(document, report.narrative.get("retraining_context", ""))
        if include_confidence:
            self._add_table(
                document,
                ["Retraining Feasibility", "Confidence Score", "Rationale"],
                [
                    [
                        report.narrative.get("retraining_feasibility", "Possible"),
                        f"{report.narrative.get('retraining_confidence_pct', 0)}%",
                        report.narrative.get("retraining_rationale", ""),
                    ]
                ],
                OxmlElement,
                qn,
                RGBColor,
                Pt,
            )
        self._add_table(
            document,
            ["Type", "Category", "Applicable Programs / Systems"],
            [
                [
                    item.get("type", ""),
                    item.get("category", ""),
                    ", ".join(item.get("applicable_programs", [])),
                ]
                for item in report.narrative.get("retraining_rows", [])
            ],
            OxmlElement,
            qn,
            RGBColor,
            Pt,
        )

        self._add_section_heading(document, "Federal & State Research and Development Credit")
        self._add_paragraph(document, report.narrative.get("rd_intro", ""))
        self._add_paragraph_break(document)
        self._add_paragraph(document, report.narrative.get("rd_examples_intro", ""))
        if include_confidence:
            self._add_table(
                document,
                ["R&D Feasibility", "Confidence Score", "Rationale"],
                [
                    [
                        report.narrative.get("rd_feasibility", "Possible"),
                        f"{report.narrative.get('rd_confidence_pct', 0)}%",
                        report.narrative.get("rd_rationale", ""),
                    ]
                ],
                OxmlElement,
                qn,
                RGBColor,
                Pt,
            )
        rd_rows = [
            [
                "R&D Activity",
                str(row.get("category", "")),
                ", ".join(row.get("activities", [])),
            ]
            for row in report.narrative.get("rd_rows", [])
        ]
        if not rd_rows:
            rd_rows = [
                ["R&D Activity", "Potential Activity", str(example)]
                for example in report.narrative.get("rd_focus_examples", [])
            ]
        self._add_table(
            document,
            ["Type", "Category", "Potential Qualifying Activities"],
            rd_rows,
            OxmlElement,
            qn,
            RGBColor,
            Pt,
        )

        self._add_section_heading(document, "Georgia Investment Tax Credit")
        self._add_paragraph(document, report.narrative.get("investment_intro", ""))
        self._add_paragraph_break(document)
        self._add_paragraph(document, report.narrative.get("investment_note", ""))
        if include_confidence:
            self._add_table(
                document,
                ["ITC Feasibility", "Confidence Score", "Rationale"],
                [
                    [
                        str(report.narrative.get("investment_status", "possible")).title(),
                        f"{report.narrative.get('investment_confidence_pct', 0)}%",
                        report.narrative.get("investment_rationale", ""),
                    ]
                ],
                OxmlElement,
                qn,
                RGBColor,
                Pt,
            )
        self._add_table(
            document,
            ["County", "Tier", "Investment Tax Credit %"],
            [
                [
                    loc.county or "-",
                    loc.ga_tier_label or "Unmapped",
                    loc.investment_tax_credit_pct or "Needs verification",
                ]
                for loc in report.locations
            ],
            OxmlElement,
            qn,
            RGBColor,
            Pt,
        )

        self._add_section_heading(document, "Cost Segregation")
        self._add_paragraph(document, report.narrative.get("costseg_intro", ""))
        self._add_paragraph_break(document)
        self._add_paragraph(document, report.narrative.get("costseg_note", ""))

        return document

    def export_report(self, report: Report, include_confidence: bool = False) -> Path:
        filename = f"{self._safe_slug(report.company_name)}_{report.id}.docx"
        output_path = self.exports_dir / filename
        document = self._build_document(report, include_confidence=include_confidence)
        document.save(output_path)
        if self.google_drive.is_enabled():
            metadata_path = self.upload_metadata_path(output_path.name)
            try:
                result = self.google_drive.upload_docx(output_path)
                if result:
                    self.google_drive.write_upload_metadata(metadata_path, result)
            except Exception as exc:
                self.google_drive.write_upload_error(metadata_path, str(exc))
        return output_path

    def upload_metadata_path(self, filename: str) -> Path:
        return self.upload_meta_dir / f"{filename}.json"

    def get_upload_metadata(self, filename: str) -> dict | None:
        return self.google_drive.read_upload_metadata(self.upload_metadata_path(filename))
